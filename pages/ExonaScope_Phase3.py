import streamlit as st
import requests
import os
from datetime import date
from io import BytesIO
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
from fpdf import FPDF
from openai import OpenAI
import hashlib

FONT_PATH = "/usr/share/fonts/truetype/msttcorefonts/Century_Schoolbook.ttf"  # Adjust this for your env!

# --- Custom CSS for streamlit preview ---
st.markdown("""
<style>
body, textarea, input, select, label, div, .stTextArea textarea, .stTextInput input, .stButton button {
    font-family: "Century Schoolbook", "Georgia", serif !important;
    font-size: 12pt !important;
}
.stMarkdown, .stTextArea textarea {
    font-family: "Century Schoolbook", "Georgia", serif !important;
    font-size: 12pt !important;
}
a { color: #2265bc !important; }
b { font-weight: 700 !important; }
</style>
""", unsafe_allow_html=True)

# ---- Jurisdictions ----
JURIS_LIST = [
    # ...[as before, full list of states/territories/federal appellate circuits]...
    ("Alabama", "alabama"), ("Alaska", "alaska"), ("Arizona", "arizona"),
    ("Arkansas", "arkansas"), ("California", "california"), ("Colorado", "colorado"),
    ("Connecticut", "connecticut"), ("Delaware", "delaware"), ("District of Columbia", "dc"),
    ("Florida", "florida"), ("Georgia", "georgia"), ("Hawaii", "hawaii"), ("Idaho", "idaho"),
    ("Illinois", "illinois"), ("Indiana", "indiana"), ("Iowa", "iowa"), ("Kansas", "kansas"),
    ("Kentucky", "kentucky"), ("Louisiana", "louisiana"), ("Maine", "maine"), ("Maryland", "maryland"),
    ("Massachusetts", "massachusetts"), ("Michigan", "michigan"), ("Minnesota", "minnesota"),
    ("Mississippi", "mississippi"), ("Missouri", "missouri"), ("Montana", "montana"), ("Nebraska", "nebraska"),
    ("Nevada", "nevada"), ("New Hampshire", "new_hampshire"), ("New Jersey", "new_jersey"),
    ("New Mexico", "new_mexico"), ("New York", "new_york"), ("North Carolina", "north_carolina"),
    ("North Dakota", "north_dakota"), ("Ohio", "ohio"), ("Oklahoma", "oklahoma"),
    ("Oregon", "oregon"), ("Pennsylvania", "pennsylvania"), ("Puerto Rico", "pr"),
    ("Rhode Island", "rhode_island"), ("South Carolina", "south_carolina"),
    ("South Dakota", "south_dakota"), ("Tennessee", "tennessee"), ("Texas", "texas"),
    ("Utah", "utah"), ("Vermont", "vermont"), ("Virginia", "virginia"),
    ("Washington", "washington"), ("West Virginia", "west_virginia"), ("Wisconsin", "wisconsin"),
    ("Wyoming", "wyoming"), ("American Samoa", "as"), ("Guam", "gu"),
    ("Northern Mariana Islands", "mp"), ("Virgin Islands", "vi"),
    ("Supreme Court of the United States", "scotus"),
    ("1st Cir. Court of Appeals", "ca1"), ("2nd Cir. Court of Appeals", "ca2"),
    ("3rd Cir. Court of Appeals", "ca3"), ("4th Cir. Court of Appeals", "ca4"),
    ("5th Cir. Court of Appeals", "ca5"), ("6th Cir. Court of Appeals", "ca6"),
    ("7th Cir. Court of Appeals", "ca7"), ("8th Cir. Court of Appeals", "ca8"),
    ("9th Cir. Court of Appeals", "ca9"), ("10th Cir. Court of Appeals", "ca10"),
    ("11th Cir. Court of Appeals", "ca11"), ("D.C. Circuit", "cadc"), ("Federal Circuit", "cafc"),
]

def bluebook_citation(case):
    if not case.get("case_name") or not case.get("citation"): return ""
    citation = f"*{case['case_name']}*, {case['citation']} ({case['court']} {case['date'][:4]})"
    if case.get("url"): return f"[{citation}]({case['url']})"
    return citation

def bluebook_citation_docx(case):
    if not case.get("case_name") or not case.get("citation"): return ""
    return f"{case['case_name']}, {case['citation']} ({case['court']} {case['date'][:4]})"

def dedup_citations(cases):
    unique = {}
    for c in cases:
        key = (c.get("citation", ""), c.get("court", ""))
        if key not in unique and c.get("case_name"):
            unique[key] = c
    return list(unique.values())

def fetch_caselaw_from_courtlistener(arg, jurisdictions, limit=4, appellate_only=False):
    """Get up-to-4 caselaw hits per jurisdiction (deduped)."""
    results = []
    for juris_code in jurisdictions:
        params = {
            "q": arg,
            "type": "o",
            "page_size": limit,
            "order_by": "-date_filed",
            "jurisdiction": juris_code,
        }
        if appellate_only:
            params["court_type"] = "A"  # "A" for appellate courts; omit for all
        params = {k: v for k, v in params.items() if v is not None}
        try:
            r = requests.get("https://www.courtlistener.com/api/rest/v3/search/", params=params, timeout=10)
            if r.status_code == 200:
                for item in r.json().get("results", []):
                    case_name = item.get("caseName") or item.get("case_name") or ""
                    citation = item.get("citation", "")
                    court = item.get("court", {}).get("name", "")
                    date_val = item.get("dateFiled", item.get("date_filed", ""))
                    url = item.get("absolute_url", "")
                    summary = item.get("plain_text", "")
                    if summary:
                        summary = summary[:350].replace("\n", " ") + ("..." if len(summary) > 340 else "")
                    results.append({
                        "case_name": case_name, "citation": citation, "court": court,
                        "date": date_val, "url": f"https://www.courtlistener.com{url}" if url else "",
                        "summary": summary
                    })
        except Exception:
            pass
    return dedup_citations(results)

def gpt_argument_and_rebuttal(section_title, arg, facts, jurisdiction_str, caselaw_md, is_suppression=True):
    api_key = os.getenv("OPENAI_API_KEY")
    client = OpenAI(api_key=api_key)
    what = 'suppression issue' if is_suppression else 'defense theory'
    prompt = f"""You are an experienced criminal defense attorney. For a confidential internal memorandum, draft a clear, highly professional legal argument for:
{what.title()}: {section_title}
Jurisdictions: {jurisdiction_str}
Facts: {facts}
Argument/Explanation: {arg}
Supporting Caselaw:
{caselaw_md}

Cite using Bluebook format (name, citation, year), and then write a short 'Counterarguments and Rebuttal' section that anticipates and responds to how the prosecution will likely attack this argument. Use real cited cases in both main argument and the rebuttal if possible. Label each section clearly.
"""
    response = client.chat.completions.create(
        model="gpt-4o",
        messages=[
            {"role": "system", "content": "You are a skilled criminal defense legal memo writer."},
            {"role": "user", "content": prompt}
        ]
    )
    return response.choices[0].message.content.strip()

def build_case_analysis_memo_docx(title, atty, case_num, date_str, facts, suppression_issues, defense_sections):
    doc = Document()
    section = doc.sections[0]
    section.left_margin = section.right_margin = Pt(72)
    section.top_margin = section.bottom_margin = Pt(72)

    # Single Header
    p = doc.add_paragraph(title)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    r = p.runs[0]
    r.font.size = Pt(16)
    r.font.name = "Century Schoolbook"
    r.bold = True
    p.paragraph_format.space_after = Pt(24)

    meta = doc.add_paragraph(f"Attorney: {atty}    Case Number: {case_num}    Date: {date_str}")
    meta.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    meta.paragraph_format.space_after = Pt(12)

    priv = doc.add_paragraph("ATTORNEY–CLIENT PRIVILEGED / WORK PRODUCT")
    priv.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    priv.runs[0].bold = True
    priv.paragraph_format.space_after = Pt(12)

    # Facts section
    head = doc.add_paragraph("SUMMARY OF PERTINENT FACTS")
    head.runs[0].bold = True
    para = doc.add_paragraph(facts)
    para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    para.paragraph_format.first_line_indent = Pt(24)
    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    para.paragraph_format.space_after = Pt(12)

    # Suppression Issues
    head = doc.add_paragraph("A. SUPPRESSION ISSUES")
    head.runs[0].bold = True
    for idx, s in enumerate(suppression_issues, 1):
        t = doc.add_paragraph(f"{idx}. {s['title']}")
        t.runs[0].bold = True
        body = doc.add_paragraph(s['argument'])
        body.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        body.paragraph_format.first_line_indent = Pt(24)
        body.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        # Caselaw
        para = doc.add_paragraph("Supporting Caselaw:")
        para.runs[0].bold = True
        if s['cases']:
            for case in s['cases']:
                txt = bluebook_citation_docx(case)
                if case.get('summary'): txt += f" — {case['summary']}"
                doc.add_paragraph(txt, style='List Bullet')
        else:
            doc.add_paragraph("No relevant caselaw found.", style='List Bullet')
        # Counter/rebuttal
        para = doc.add_paragraph("Counterarguments and Rebuttal:")
        para.runs[0].bold = True
        rebut = doc.add_paragraph(s.get('rebuttal', ''))
        rebut.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        rebut.paragraph_format.first_line_indent = Pt(24)
        rebut.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    # Defenses
    head = doc.add_paragraph("B. POTENTIAL DEFENSES")
    head.runs[0].bold = True
    for idx, d in enumerate(defense_sections, 1):
        t = doc.add_paragraph(f"{idx}. {d['title']}")
        t.runs[0].bold = True
        body = doc.add_paragraph(d['argument'])
        body.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        body.paragraph_format.first_line_indent = Pt(24)
        body.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        # Caselaw
        para = doc.add_paragraph("Supporting Caselaw:")
        para.runs[0].bold = True
        if d['cases']:
            for case in d['cases']:
                txt = bluebook_citation_docx(case)
                if case.get('summary'): txt += f" — {case['summary']}"
                doc.add_paragraph(txt, style='List Bullet')
        else:
            doc.add_paragraph("No relevant caselaw found.", style='List Bullet')
        para = doc.add_paragraph("Counterarguments and Rebuttal:")
        para.runs[0].bold = True
        rebut = doc.add_paragraph(d.get('rebuttal', ''))
        rebut.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        rebut.paragraph_format.first_line_indent = Pt(24)
        rebut.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    c = doc.add_paragraph("CONCLUSION")
    c.runs[0].bold = True
    concl_text = "This memorandum is for internal defense team review only and is not intended for filing without attorney revision.\n"
    concl = doc.add_paragraph(concl_text)
    concl.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # FONT consistency on every run
    for p in doc.paragraphs:
        for r in p.runs:
            r.font.name = "Century Schoolbook"
            r.font.size = Pt(12)
    return doc

def text_to_pdf(text):
    pdf = FPDF()
    pdf.add_page()
    try:
        pdf.add_font('CSchoolbook', '', FONT_PATH, uni=True)
        pdf.set_font("CSchoolbook", '', 12)
    except Exception as e:
        print(f"Error occurred: {e}")
        pdf.set_font("Arial", size=12)
    safe_text = text.replace('\u2013', '-')
    for line in safe_text.split('\n'):
        pdf.multi_cell(0, 10, line if isinstance(line, str) else str(line))
        pdf_string = pdf.output(dest="S")
        pdf_bytes = pdf_string  # Direct use of bytearray
        pdf_output = BytesIO(pdf_bytes)


        # Capture the output of the PDF
        output = pdf.output(dest='S')  # This directly gives a bytearray
        pdf_output.write(output)

        pdf_output.write(output)
        pdf_output.seek(0)
        return pdf_output

def content_hash(title, argument):
    return hashlib.md5((title + argument).encode()).hexdigest()

# ==== Streamlit UI ====
st.title("ExonaScope Phase 3 – Case Analysis Memorandum")

# --- Editable Inputs: Attorney, Case Number, Date ---
attorney_name = st.text_input("Attorney Name", value=st.session_state.get("attorney_name", ""), key="attorney_name")
case_number = st.text_input("Case Number", value=st.session_state.get("case_number", ""), key="case_number")
today_date = date.today().strftime("%B %d, %Y")

# --- Editable Facts ---
memo_facts = st.text_area("Summary of Pertinent Facts", value=st.session_state.get("motion_facts", ""), key="summary_facts")

# --- Load Phase 2 Data ---
phase2_issues = st.session_state.get("phase2_issues", []) or []
phase2_defenses = st.session_state.get("phase2_defenses", []) or []

# --- Jurisdictions, Appellate Only ---
juris_selected = st.multiselect("Select Jurisdictions:",
    options=JURIS_LIST,
    default=[("Supreme Court of the United States", "scotus")],
    format_func=lambda x: x[0]
)
juris_codes = [code for desc, code in juris_selected]
juris_label = ", ".join(desc for desc, code in juris_selected)
appellate_only = st.checkbox("Appellate Cases Only", value=True)

# --- Session State for Boxes ---
if "issue_boxes" not in st.session_state:
    st.session_state.issue_boxes = [
        {"title": i["title"], "argument": i.get("explanation", "")} for i in phase2_issues]
if "defense_boxes" not in st.session_state:
    st.session_state.defense_boxes = [
        {"title": d["title"], "argument": d.get("explanation", "")} for d in phase2_defenses]

# --- UI for Issues/Defenses --
if st.button("Add Custom Suppression Issue"):
    st.session_state.issue_boxes.append({"title": "", "argument": ""})
if st.button("Add Custom Defense"):
    st.session_state.defense_boxes.append({"title": "", "argument": ""})

st.subheader("Suppression Issues (edit or add as needed)")
issue_args = []
for idx, box in enumerate(st.session_state.issue_boxes):
    c1, c2 = st.columns([1,4])
    box["title"] = c1.text_input(f"Issue {idx+1} Short Title", value=box["title"], key=f"issue_title_{idx}")
    box["argument"] = c2.text_area(f"Issue {idx+1} Argument/Explanation", value=box["argument"], key=f"issue_argument_{idx}")
    issue_args.append(box)

st.subheader("Defense Theories (edit or add as needed)")
defense_args = []
for idx, box in enumerate(st.session_state.defense_boxes):
    c1, c2 = st.columns([1,4])
    box["title"] = c1.text_input(f"Defense {idx+1} Short Title", value=box["title"], key=f"defense_title_{idx}")
    box["argument"] = c2.text_area(f"Defense {idx+1} Argument/Explanation", value=box["argument"], key=f"defense_argument_{idx}")
    defense_args.append(box)

# --- Require Key Fields to Export ---
allow_export = attorney_name.strip() and case_number.strip()
if not allow_export:
    st.warning("Please enter both an Attorney Name and Case Number.")

# --- Run Caselaw Search & Generate Memo ---
if st.button("Run Caselaw Search & Generate Memo") and allow_export:
    suppression_sections = []
    defense_sections = []
    # Dirty tracking for suppression
    for idx, issue in enumerate(issue_args):
        cur_hash = content_hash(issue["title"], issue["argument"])
        hash_key, res_key = f"issue_hash_{idx}", f"issue_result_{idx}"
        if st.session_state.get(hash_key) != cur_hash or not st.session_state.get(res_key):
            search_arg = f"{issue['title']} {issue['argument']}".strip()
            cases = fetch_caselaw_from_courtlistener(search_arg, juris_codes, limit=4, appellate_only=appellate_only)
            case_md_list = []
            for c in cases:
                bb = bluebook_citation(c)
                if c.get('summary'):
                    bb += f" — {c['summary']}"
                case_md_list.append(bb)
            memo_full = gpt_argument_and_rebuttal(issue['title'], issue['argument'], memo_facts, juris_label, "\n".join(case_md_list), True)
            main, rebuttal = memo_full, ""
            if "Counterarguments and Rebuttal:" in memo_full:
                parts = memo_full.split("Counterarguments and Rebuttal:")
                main = parts[0].strip()
                rebuttal = parts[1].strip() if len(parts) > 1 else ""
            st.session_state[hash_key], st.session_state[res_key] = cur_hash, {"title": issue["title"], "argument": main, "cases": cases, "rebuttal": rebuttal}
        section = st.session_state[res_key]
        suppression_sections.append(section)

    # Dirty tracking for defenses
    for idx, defense in enumerate(defense_args):
        cur_hash = content_hash(defense["title"], defense["argument"])
        hash_key, res_key = f"defense_hash_{idx}", f"defense_result_{idx}"
        if st.session_state.get(hash_key) != cur_hash or not st.session_state.get(res_key):
            search_arg = f"{defense['title']} {defense['argument']}".strip()
            cases = fetch_caselaw_from_courtlistener(search_arg, juris_codes, limit=4, appellate_only=appellate_only)
            case_md_list = []
            for c in cases:
                bb = bluebook_citation(c)
                if c.get('summary'):
                    bb += f" — {c['summary']}"
                case_md_list.append(bb)
            memo_full = gpt_argument_and_rebuttal(defense['title'], defense['argument'], memo_facts, juris_label, "\n".join(case_md_list), False)
            main, rebuttal = memo_full, ""
            if "Counterarguments and Rebuttal:" in memo_full:
                parts = memo_full.split("Counterarguments and Rebuttal:")
                main = parts[0].strip()
                rebuttal = parts[1].strip() if len(parts) > 1 else ""
            st.session_state[hash_key], st.session_state[res_key] = cur_hash, {"title": defense["title"], "argument": main, "cases": cases, "rebuttal": rebuttal}
        section = st.session_state[res_key]
        defense_sections.append(section)

    # --- Memo Preview ---
    memo_lines = []
    memo_lines.append("<div style='text-align: center; font-weight: bold; font-size:20px;'>CASE ANALYSIS MEMORANDUM</div>")
    memo_lines.append(f"<b>Attorney:</b> {attorney_name} &nbsp;&nbsp;&nbsp; <b>Case Number:</b> {case_number} &nbsp;&nbsp;&nbsp; <b>Date:</b> {today_date}<br>")
    memo_lines.append("<b>ATTORNEY–CLIENT PRIVILEGED / WORK PRODUCT</b><br>")
    memo_lines.append(f"<b>SUMMARY OF PERTINENT FACTS</b><br>{memo_facts}<br>")
    memo_lines.append("<b>A. SUPPRESSION ISSUES</b>")
    for idx, s in enumerate(suppression_sections, 1):
        memo_lines.append(f"<b>{idx}. {s['title']}</b><br>{s['argument']}<br><b>Supporting Caselaw:</b>")
        if s['cases']:
            for c in s['cases']:
                bb = bluebook_citation(c)
                if c.get('summary'):
                    bb += f" — {c['summary']}"
                memo_lines.append(f"&nbsp;&nbsp;• {bb}")
        else:
            memo_lines.append("&nbsp;&nbsp;<i>No relevant caselaw found for this argument in the selected jurisdictions.</i>")
        if s['rebuttal']:
            memo_lines.append(f"<b>Counterarguments and Rebuttal:</b> {s['rebuttal']}")
        memo_lines.append("")
    memo_lines.append("<b>B. POTENTIAL DEFENSES</b>")
    for idx, d in enumerate(defense_sections, 1):
        memo_lines.append(f"<b>{idx}. {d['title']}</b><br>{d['argument']}<br><b>Supporting Caselaw:</b>")
        if d['cases']:
            for c in d['cases']:
                bb = bluebook_citation(c)
                if c.get('summary'):
                    bb += f" — {c['summary']}"
                memo_lines.append(f"&nbsp;&nbsp;• {bb}")
        else:
            memo_lines.append("&nbsp;&nbsp;<i>No relevant caselaw found for this argument in the selected jurisdictions.</i>")
        if d['rebuttal']:
            memo_lines.append(f"<b>Counterarguments and Rebuttal:</b> {d['rebuttal']}")
        memo_lines.append("")
    memo_lines.extend([
        "",
        "<b>CONCLUSION</b>",
        "This memorandum is for internal defense team review only and is not intended for filing without attorney revision.",
        "Respectfully submitted,",
        "<br><br>____________________________<br>Attorney for the Defense"
    ])
    st.markdown('<br>'.join(memo_lines), unsafe_allow_html=True)

    # --- DOCX export ---
    doc_obj = build_case_analysis_memo_docx(
        "CASE ANALYSIS MEMORANDUM",
        attorney_name, case_number, today_date, memo_facts, suppression_sections, defense_sections
    )
    docx_bytes = BytesIO(); doc_obj.save(docx_bytes); docx_bytes.seek(0)
    st.download_button("📥 Download Memo (.docx)", data=docx_bytes.getvalue(),
        file_name="Case_Analysis_Memorandum.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

    # --- PDF Export (plain but readable) ---
    plain_preview = '\n'.join([c.replace('<br>', '\n').replace('&nbsp;', ' ') for c in memo_lines])
    pdf_bytes = text_to_pdf(plain_preview)
    st.download_button("📄 Download Memo as PDF",
        data=pdf_bytes,
        file_name="Case_Analysis_Memorandum.pdf",
        mime="application/pdf"
    )

if st.checkbox("🪵 Show Session State"):
    st.json(dict(st.session_state))
