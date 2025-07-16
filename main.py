import streamlit as st
import os
import tempfile
import requests
import time
from docx import Document
from io import BytesIO
import fitz
import pytesseract
from pdf2image import convert_from_bytes
from PIL import Image
import subprocess
import mimetypes
from moviepy.video.io.VideoFileClip import VideoFileClip

# --- CONFIG ---
ASSEMBLYAI_API_KEY = os.environ.get("ASSEMBLYAI_API_KEY")
OPENAI_API_KEY = os.environ.get("OPENAI_API_KEY")

if not ASSEMBLYAI_API_KEY:
    st.error("Please set your ASSEMBLYAI_API_KEY in your environment variables.")
    st.stop()

# --- Utility Functions ---

def transcribe_with_assemblyai_from_path(filepath):
    headers = {"authorization": ASSEMBLYAI_API_KEY}
    with open(filepath, "rb") as f:
        upload_response = requests.post("https://api.assemblyai.com/v2/upload", headers=headers, files={"file": f})
    if upload_response.status_code != 200:
        return f"[Upload Error: {upload_response.text}]"
    upload_url = upload_response.json()["upload_url"]
    transcript_response = requests.post("https://api.assemblyai.com/v2/transcript", headers=headers, json={"audio_url": upload_url})
    if transcript_response.status_code != 200:
        return f"[Start Error: {transcript_response.text}]"
    transcript_id = transcript_response.json()["id"]
    for _ in range(60):
        poll_response = requests.get(f"https://api.assemblyai.com/v2/transcript/{transcript_id}", headers=headers)
        status = poll_response.json()["status"]
        if status == "completed":
            return poll_response.json()["text"]
        elif status == "error":
            return f"[Transcription Error: {poll_response.json()['error']}]"
        time.sleep(3)
    return "[Timeout waiting for transcription]"

def has_audio_track(video_path):
    cmd = ["ffprobe", "-v", "error", "-select_streams", "a", "-show_entries", "stream=index", "-of", "csv=p=0", video_path]
    result = subprocess.run(cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
    return result.stdout.strip() != b''

def extract_audio_from_video(video_file_path, audio_ext=".wav"):
    with tempfile.NamedTemporaryFile(delete=False, suffix=audio_ext) as temp_audio:
        audio_path = temp_audio.name
    try:
        with VideoFileClip(video_file_path) as video:
            if video.audio is None:
                raise Exception("No audio stream found in the video.")
            video.audio.write_audiofile(audio_path, codec='pcm_s16le', fps=16000, nbytes=2)
        if not os.path.exists(audio_path):
            raise FileNotFoundError(f"Audio file not found: {audio_path}")
        return audio_path
    except Exception as e:
        return f"[Audio extraction failed: {e}]"

def fallback_extract_with_ffmpeg(video_file_path, output_audio_path):
    command = ["ffmpeg", "-y", "-i", video_file_path, "-vn", "-acodec", "pcm_s16le", "-ar", "16000", "-ac", "1", output_audio_path]
    try:
        subprocess.run(command, check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
        return output_audio_path
    except subprocess.CalledProcessError as e:
        return f"[FFmpeg fallback failed: {e.stderr.decode()}]"

def reencode_audio_to_pcm_wav(input_path):
    reencoded_path = input_path.replace(".wav", "_fixed.wav")
    command = ["ffmpeg", "-y", "-i", input_path, "-acodec", "pcm_s16le", "-ar", "16000", "-ac", "1", reencoded_path]
    try:
        subprocess.run(command, check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
        return reencoded_path
    except subprocess.CalledProcessError as e:
        return f"[Re-encode failed: {e.stderr.decode()}]"

def convert_wav_to_mp3(input_path):
    mp3_path = input_path.replace(".wav", ".mp3")
    command = ["ffmpeg", "-y", "-i", input_path, "-acodec", "libmp3lame", "-ar", "16000", "-ac", "1", mp3_path]
    try:
        subprocess.run(command, check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
        return mp3_path
    except subprocess.CalledProcessError as e:
        return f"[MP3 conversion failed: {e.stderr.decode()}]"

def parse_pdf_text(file):
    file.seek(0)
    doc = fitz.open(stream=file.read(), filetype="pdf")
    return "\n".join([page.get_text() for page in doc if page.get_text()])

def run_ocr_on_pdf(file):
    file.seek(0)
    images = convert_from_bytes(file.read(), dpi=300)
    text = ""
    for img in images:
        text += pytesseract.image_to_string(img, config="--psm 6")
    return text

def parse_docx(file):
    return "\n".join([p.text for p in Document(file).paragraphs])

def extract_text_from_file(uploaded_file):
    parsed = ""
    temp_paths = []
    if uploaded_file.type.startswith("video/") or uploaded_file.name.lower().endswith((".mp4", ".avi", ".mkv", ".mov")):
        st.info("Extracting audio from video...")
        with tempfile.NamedTemporaryFile(delete=False, suffix=".mp4") as temp_video:
            video_bytes = uploaded_file.read()
            temp_video.write(video_bytes)
            temp_video.flush()
            os.fsync(temp_video.fileno())
            temp_video_path = temp_video.name
            temp_paths.append(temp_video_path)

        if not has_audio_track(temp_video_path):
            return "[Error: No audio stream detected.]"

        audio_path = extract_audio_from_video(temp_video_path)
        if isinstance(audio_path, str) and audio_path.startswith("[Audio extraction failed"):
            with tempfile.NamedTemporaryFile(delete=False, suffix=".wav") as temp_fallback:
                fallback_path = temp_fallback.name
            audio_path = fallback_extract_with_ffmpeg(temp_video_path, fallback_path)
            temp_paths.append(fallback_path)
        temp_paths.append(audio_path)

        if not os.path.exists(audio_path) or os.path.getsize(audio_path) == 0:
            return "[Transcription Error: Extracted audio file is missing or empty.]"

        fixed_wav_path = reencode_audio_to_pcm_wav(audio_path)
        mp3_path = convert_wav_to_mp3(fixed_wav_path)
        temp_paths.extend([fixed_wav_path, mp3_path])

        mime, _ = mimetypes.guess_type(mp3_path)
        st.audio(mp3_path)
        st.write(f"🧪 Uploading: {mp3_path} | MIME: {mime} | Size: {os.path.getsize(mp3_path)} bytes")

        st.info("Transcribing re-encoded MP3 audio...")
        parsed = transcribe_with_assemblyai_from_path(mp3_path)

        for path in temp_paths:
            if os.path.exists(path):
                os.remove(path)

    elif "pdf" in uploaded_file.type:
        parsed = parse_pdf_text(uploaded_file)
        if not parsed.strip():
            st.info("No embedded text, running OCR...")
            parsed = run_ocr_on_pdf(uploaded_file)
    elif "word" in uploaded_file.type or uploaded_file.name.endswith(".docx"):
        parsed = parse_docx(uploaded_file)
    elif "audio" in uploaded_file.type or uploaded_file.type.startswith("audio/"):
        uploaded_file.seek(0)
        ext = os.path.splitext(uploaded_file.name)[1].lower()
        with tempfile.NamedTemporaryFile(delete=False, suffix=ext) as tmp:
            tmp.write(uploaded_file.read())
            tmp.flush()
            tmp_path = tmp.name
        with st.spinner("Transcribing audio..."):
            parsed = transcribe_with_assemblyai_from_path(tmp_path)
        os.remove(tmp_path)
    else:
        parsed = "[Unsupported file type]"
    return parsed

def save_docx(text, filename="output.docx"):
    docx_file = BytesIO()
    doc = Document()
    doc.add_paragraph(text)
    doc.save(docx_file)
    docx_file.seek(0)
    return docx_file

def extract_facts_with_gpt_chunked(full_text, case_name, case_number, chunk_size=4000):
    if not OPENAI_API_KEY:
        return "[OpenAI API key not set. Cannot extract facts.]"
    import openai
    client = openai.OpenAI(api_key=OPENAI_API_KEY)
    chunks = [full_text[i:i+chunk_size] for i in range(0, len(full_text), chunk_size)]
    all_facts = []
    for idx, chunk in enumerate(chunks):
        prompt = f"""Using only the exact facts from the following material — without combining, summarizing, or paraphrasing — extract every individual event and action exactly as written, in strict chronological order.

CASE NAME: {case_name}
CASE NUMBER: {case_number}
SOURCE MATERIAL (PART {idx+1} of {len(chunks)}):

{chunk}
"""
        with st.spinner(f"Extracting facts from chunk {idx+1} of {len(chunks)}..."):
            try:
                response = client.chat.completions.create(
                    model="gpt-4o",
                    messages=[
                        {"role": "system", "content": "You extract and present only the original facts in strict chronological order for legal suppression review. Do not enhance."},
                        {"role": "user", "content": prompt}
                    ],
                    temperature=0.0
                )
                facts = response.choices[0].message.content.strip()
                all_facts.append(facts)
            except Exception as e:
                all_facts.append(f"[GPT Error in chunk {idx+1}: {e}]")
    return "\n\n".join(all_facts)

# --- UI ---
st.title("ExonaScope Phase 1 – Upload, Transcribe, Extract Facts")
case_name = st.text_input("Case Name")
case_number = st.text_input("Case Number")
uploaded_files = st.file_uploader(
    "Upload PDFs, DOCX, audio, or video files",
    type=["pdf", "docx", "mp3", "wav", "m4a", "mp4", "avi", "mkv", "mov"],
    accept_multiple_files=True
)

parsed_segments = []
if uploaded_files:
    st.subheader("📄 Parsed Preview")
    for idx, uploaded_file in enumerate(uploaded_files):
        st.write(f"**File:** {uploaded_file.name}")
        try:
            parsed = extract_text_from_file(uploaded_file)
            if parsed.strip():
                parsed_segments.append(f"[{uploaded_file.name}]\n{parsed}")
                with st.expander(f"Preview: {uploaded_file.name}"):
                    st.text(parsed[:2000])
                if (
                    ("audio" in uploaded_file.type or uploaded_file.type.startswith("audio/")) or
                    (uploaded_file.type.startswith("video/") or uploaded_file.name.lower().endswith((".mp4", ".avi", ".mkv", ".mov")))
                ) and not parsed.startswith("["):
                    docx_file = save_docx(parsed, filename=f"{uploaded_file.name}_transcript.docx")
                    st.download_button(
                        f"Download Transcript ({uploaded_file.name})",
                        docx_file,
                        file_name=f"{uploaded_file.name}_transcript.docx",
                        key=f"download_transcript_{idx}"
                    )
            else:
                st.warning(f"⚠️ Nothing extractable from: {uploaded_file.name}")
        except Exception as e:
            st.error(f"❌ Error processing {uploaded_file.name}: {e}")

# --- Fact Extraction and Editing ---
if parsed_segments:
    if st.button("🧠 Generate Chronological Facts (GPT-4o)", key="generate_facts"):
        full_text = "\n\n".join(parsed_segments)
        facts = extract_facts_with_gpt_chunked(full_text, case_name, case_number, chunk_size=4000)
        if facts and not facts.startswith("["):
            st.success("Fact extraction complete!")
            st.session_state["phase2_facts"] = facts
        else:
            st.error(facts)

    # Show editable facts if available
    if "phase2_facts" in st.session_state and st.session_state["phase2_facts"]:
        facts_editable = st.text_area(
            "Facts (edit before continuing to Phase 2):",
            value=st.session_state["phase2_facts"],
            height=300,
            key="facts_editable"
        )
        st.session_state["phase2_facts"] = facts_editable
        docx_file = save_docx(facts_editable, filename="facts.docx")
        st.download_button(
            "Download Facts (.docx)",
            docx_file,
            file_name="facts.docx",
            key="download_facts"
        )

# --- Handoff to Phase 2 ---
if st.button("Continue to Legal Analysis in Phase 2", key="continue_phase2"):
    st.session_state["case_name"] = case_name
    st.session_state["case_number"] = case_number
    # Ensure facts are present
    st.session_state["phase2_facts"] = st.session_state.get("phase2_facts", "")
    st.switch_page("pages/ExonaScope_Phase2.py")  # Use the correct path to your Phase 2 script

# --- Debugging: Show Session State ---
if st.checkbox("Show Session State (Debug)", key="show_session_state"):
    st.write(dict(st.session_state))

